"""packaging.py - render a pack (md) into html/pdf/docx/zip. Uses pandoc when available, pure-python otherwise."""
import os, sys, zipfile, subprocess, json

def md_to_html(md_path, html_path, title):
    html = open(md_path).read()
    try:
        import markdown
        body = markdown.markdown(html, extensions=["tables", "fenced_code"])
    except ImportError:
        body = "<pre>" + html.replace("<", "&lt;") + "</pre>"
    css = "body{font-family:system-ui,Segoe UI,Roboto,sans-serif;max-width:800px;margin:2rem auto;line-height:1.6}h1,h2{color:#111}code{background:#f4f4f4;padding:2px 4px}pre{background:#0d1117;color:#e6edf3;padding:1rem;border-radius:8px}"
    out = f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title><style>{css}</style></head><body>{body}</body></html>"
    open(html_path, "w").write(out)

def html_to_pdf(html_path, pdf_path):
    from weasyprint import HTML
    HTML(filename=html_path).write_pdf(pdf_path)

def html_to_docx(html_path, docx_path, title):
    from docx import Document
    import re
    doc = Document(); doc.add_heading(title, 0)
    raw = open(html_path).read().replace("\n", " ")
    # crude html->docx: headings + paragraphs + code blocks
    for m in re.findall(r"<(h[12])[^>]*>(.*?)</\1>|<p>(.*?)</p>|<pre[^>]*>(.*?)</pre>", raw, re.S):
        h, _, p, code = (m[0], m[1], m[2], m[3])
        if h: doc.add_heading(re.sub(r"<[^>]+>", "", m[1]).strip(), level=int(h[1]))
        elif p: doc.add_paragraph(re.sub(r"<[^>]+>", "", p).strip())
        elif code: doc.add_paragraph(code, style="No Spacing")
    doc.save(docx_path)

def pack_all(base_dir, title, out_dir):
    md = os.path.join(base_dir, "pack.md"); html = os.path.join(base_dir, "pack.html")
    os.makedirs(out_dir, exist_ok=True)
    pdf = os.path.join(out_dir, f"{title}.pdf"); docx = os.path.join(out_dir, f"{title}.docx")
    z = os.path.join(out_dir, f"{title}.zip")
    # 1. html always from markdown
    md_to_html(md, html, title)
    # 2. pdf: pandoc if available else weasyprint
    try:
        subprocess.run(["pandoc", md, "-o", pdf, "--pdf-engine=weasyprint"], check=True, capture_output=True)
    except Exception:
        html_to_pdf(html, pdf)
    html_to_docx(html, docx, title)
    with zipfile.ZipFile(z, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in ("pack.md", "pack.html"):
            zf.write(os.path.join(base_dir, f), f)
        zf.write(pdf, os.path.basename(pdf)); zf.write(docx, os.path.basename(docx))
    for f in (pdf, docx, z):
        print(f"{f}: {os.path.getsize(f)} bytes")
    return {"pdf": pdf, "docx": docx, "zip": z, "html": html}

if __name__ == "__main__":
    pack_all(sys.argv[1], sys.argv[2], sys.argv[3])
